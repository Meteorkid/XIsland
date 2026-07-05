#!/usr/bin/env python3
"""
验证每页是否正好50行代码
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def verify_lines_per_page():
    """验证每页的代码行数"""
    print("验证每页代码行数...")

    doc = Document(OUTPUT_FILE)

    # 找到所有分页符的位置（通过检查run中的break）
    page_break_positions = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            # 检查run中是否有分页符
            if run._element.xml.find('w:br') != -1 and 'type="page"' in run._element.xml:
                page_break_positions.append(i)
                break

    print(f"找到 {len(page_break_positions)} 个分页符")

    # 统计每页的代码行数
    page_line_counts = []
    code_line_count = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 检查是否是代码行（以数字开头）
        if text and text[0].isdigit() and '  ' in text:
            code_line_count += 1

        # 检查是否到达分页符
        if i in page_break_positions:
            page_line_counts.append(code_line_count)
            code_line_count = 0

    # 添加最后一页的代码行数
    if code_line_count > 0:
        page_line_counts.append(code_line_count)

    print(f"\n=== 每页代码行数统计 ===")
    print(f"总页数: {len(page_line_counts)}")

    # 检查每页是否正好50行
    all_correct = True
    incorrect_pages = []

    for i, count in enumerate(page_line_counts):
        page_num = i + 1
        if count != 50:
            all_correct = False
            incorrect_pages.append((page_num, count))
            print(f"第 {page_num} 页: {count} 行 ❌ (应为50行)")
        else:
            print(f"第 {page_num} 页: {count} 行 ✓")

    # 总结
    print(f"\n=== 验证结果 ===")
    if all_correct:
        print(f"✅ 所有页面都是正好50行代码！")
        return True
    else:
        print(f"❌ 以下页面不是50行代码:")
        for page_num, count in incorrect_pages:
            print(f"  - 第 {page_num} 页: {count} 行")
        return False


def main():
    """主函数"""
    try:
        success = verify_lines_per_page()
        return 0 if success else 1
    except Exception as e:
        print(f"\n❌ 验证失败: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())
