#!/usr/bin/env python3
"""
检查分页符和空行的关系
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def check_page_breaks():
    """检查分页符"""
    print("检查分页符和空行的关系...")

    doc = Document(OUTPUT_FILE)

    # 找到所有分页符的位置
    page_break_positions = []
    for i, para in enumerate(doc.paragraphs):
        if para.paragraph_format.page_break_before:
            page_break_positions.append(i)

    print(f"找到 {len(page_break_positions)} 个分页符")

    # 检查分页符前后的内容
    print(f"\n=== 分页符位置分析 ===")
    for i, pos in enumerate(page_break_positions[:5]):  # 只显示前5个
        print(f"\n分页符 {i+1}: 段落 {pos}")

        # 检查分页符前的段落
        if pos > 0:
            prev_text = doc.paragraphs[pos-1].text.strip()
            if prev_text:
                print(f"  前一段落: {prev_text[:50]}...")
            else:
                print(f"  前一段落: (空)")

        # 检查分页符后的段落
        if pos < len(doc.paragraphs) - 1:
            next_text = doc.paragraphs[pos+1].text.strip()
            if next_text:
                print(f"  后一段落: {next_text[:50]}...")
            else:
                print(f"  后一段落: (空)")

    # 检查空行是否在分页符位置
    print(f"\n=== 空行与分页符关系 ===")
    empty_lines = []
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            empty_lines.append(i)

    print(f"空行总数: {len(empty_lines)}")

    # 检查空行是否在分页符位置
    empty_at_page_break = [pos for pos in empty_lines if pos in page_break_positions]
    print(f"空行在分页符位置: {len(empty_at_page_break)} 个")

    # 检查空行是否在分页符前后
    empty_near_page_break = []
    for pos in empty_lines:
        for pb_pos in page_break_positions:
            if abs(pos - pb_pos) <= 1:
                empty_near_page_break.append(pos)
                break

    print(f"空行在分页符附近: {len(empty_near_page_break)} 个")


if __name__ == "__main__":
    check_page_breaks()
