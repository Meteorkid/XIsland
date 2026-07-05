#!/usr/bin/env python3
"""
严格验证源代码文档
检查：
1. 每页正好50行代码
2. 总共60页
3. 没有空行
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def verify_document():
    """严格验证文档"""
    print("开始严格验证文档...")

    if not os.path.exists(OUTPUT_FILE):
        print(f"❌ 文件不存在: {OUTPUT_FILE}")
        return False

    # 读取文档
    doc = Document(OUTPUT_FILE)

    # 找到所有分页符的位置
    page_break_positions = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run._element.xml.find('w:br') != -1 and 'type="page"' in run._element.xml:
                page_break_positions.append(i)
                break

    print(f"找到 {len(page_break_positions)} 个分页符")

    # 统计每页的代码行数
    page_line_counts = []
    code_line_count = 0
    empty_line_count = 0
    current_page = 1

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 检查是否是代码行（以数字开头）
        if text and text[0].isdigit() and '  ' in text:
            code_line_count += 1
        elif not text:
            # 检查是否是分页符段落（只有分页符，没有文本）
            is_page_break_para = False
            for run in para.runs:
                if run._element.xml.find('w:br') != -1 and 'type="page"' in run._element.xml:
                    is_page_break_para = True
                    break

            if is_page_break_para:
                # 这是分页符段落，不计入空行
                pass
            else:
                empty_line_count += 1

        # 检查是否到达分页符
        if i in page_break_positions:
            page_line_counts.append(code_line_count)
            code_line_count = 0
            current_page += 1

    # 添加最后一页的代码行数
    if code_line_count > 0:
        page_line_counts.append(code_line_count)

    print(f"\n=== 验证结果 ===")

    # 1. 检查总页数
    total_pages = len(page_line_counts)
    print(f"总页数: {total_pages}")

    # 2. 检查空行
    print(f"空行数: {empty_line_count}")

    # 3. 检查每页代码行数（排除标题页和分隔页）
    code_pages = [count for count in page_line_counts if count > 0]
    non_code_pages = [count for count in page_line_counts if count == 0]

    print(f"代码页数: {len(code_pages)}")
    print(f"非代码页数: {len(non_code_pages)} (标题页/分隔页)")

    all_correct = True
    incorrect_pages = []

    for i, count in enumerate(page_line_counts):
        page_num = i + 1
        if count == 0:
            # 非代码页（标题页或分隔页），跳过
            continue
        elif count != 50:
            all_correct = False
            incorrect_pages.append((page_num, count))

    # 4. 检查总代码行数
    total_code_lines = sum(page_line_counts)
    print(f"总代码行数: {total_code_lines}")

    # 5. 检查文件大小
    file_size = os.path.getsize(OUTPUT_FILE) / 1024
    print(f"文件大小: {file_size:.1f} KB")

    # 总结
    print(f"\n=== 最终结论 ===")

    if all_correct and len(code_pages) == 60 and empty_line_count == 0 and total_code_lines == 3000:
        print(f"✅ 文档完全符合软著申请要求！")
        print(f"  - 每页正好50行代码 ✓")
        print(f"  - 总共60页代码 ✓")
        print(f"  - 无空行 ✓")
        print(f"  - 总代码行数3000行 ✓")
        print(f"  - 标题页和分隔页正常 ✓")
        return True
    else:
        print(f"❌ 文档未完全符合要求：")
        if not all_correct:
            print(f"  - 以下页面不是50行代码:")
            for page_num, count in incorrect_pages:
                print(f"    * 第 {page_num} 页: {count} 行")
        if len(code_pages) != 60:
            print(f"  - 代码页数不是60页 (实际: {len(code_pages)})")
        if empty_line_count != 0:
            print(f"  - 有空行 (实际: {empty_line_count} 行)")
        if total_code_lines != 3000:
            print(f"  - 总代码行数不是3000行 (实际: {total_code_lines})")
        return False


def main():
    """主函数"""
    try:
        success = verify_document()
        return 0 if success else 1
    except Exception as e:
        print(f"\n❌ 验证失败: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())
