#!/usr/bin/env python3
"""
检查文档结构
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def check_structure():
    """检查文档结构"""
    print("检查文档结构...")

    doc = Document(OUTPUT_FILE)

    # 统计信息
    total_paragraphs = len(doc.paragraphs)
    code_lines = 0
    empty_lines = 0
    title_lines = 0
    separator_lines = 0

    # 检查前20个段落
    print("\n=== 前20个段落 ===")
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if text:
            # 判断段落类型
            if text.startswith("X Island"):
                print(f"[{i}] 标题: {text[:50]}...")
                title_lines += 1
            elif text.startswith("软件著作权"):
                print(f"[{i}] 副标题: {text[:50]}...")
                title_lines += 1
            elif text[0].isdigit() and '  ' in text:
                print(f"[{i}] 代码行: {text[:50]}...")
                code_lines += 1
            elif text.startswith("// File:"):
                print(f"[{i}] 文件注释: {text[:50]}...")
                code_lines += 1
            else:
                print(f"[{i}] 其他: {text[:50]}...")
                separator_lines += 1
        else:
            print(f"[{i}] 空行")
            empty_lines += 1

    # 检查分页符
    print("\n=== 检查分页符 ===")
    page_break_count = 0
    for i, para in enumerate(doc.paragraphs):
        # 检查段落格式
        if para.paragraph_format.page_break_before:
            page_break_count += 1
            if page_break_count <= 5:  # 只显示前5个
                print(f"段落 {i}: 有分页符")

    print(f"\n总共发现 {page_break_count} 个分页符")

    # 检查代码行
    print("\n=== 检查代码行 ===")
    code_count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and text[0].isdigit() and '  ' in text:
            code_count += 1
            if code_count <= 3:  # 只显示前3行
                print(f"代码行 {code_count}: {text[:60]}...")

    print(f"\n总共发现 {code_count} 行代码")


if __name__ == "__main__":
    check_structure()
