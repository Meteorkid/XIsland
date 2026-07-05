#!/usr/bin/env python3
"""
生成软著申请源代码文档（最终版）
严格要求：
1. 每页正好50行代码
2. 总共60页（前30页+后30页）
3. 没有空行
4. 严格验证
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 配置
SOFTWARE_NAME = "X Island"
VERSION = "V1.0.0"
OUTPUT_DIR = "/Users/meteor/github/xisland/软著申请材料"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "XIsland_源代码文档.docx")
SOURCE_FILE = os.path.join(OUTPUT_DIR, "source_code_full.txt")

# 页面设置
PAGE_WIDTH = Cm(21)  # A4纸宽度
PAGE_HEIGHT = Cm(29.7)  # A4纸高度
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.5)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)

# 代码行数配置
LINES_PER_PAGE = 50
PAGES_FRONT = 30
PAGES_BACK = 30
TOTAL_PAGES = PAGES_FRONT + PAGES_BACK
TOTAL_LINES = LINES_PER_PAGE * TOTAL_PAGES  # 3000行

# 字体配置
CODE_FONT_NAME = "Courier New"  # 等宽字体
CODE_FONT_SIZE = Pt(10.5)  # 五号字
HEADER_FONT_NAME = "宋体"
HEADER_FONT_SIZE = Pt(9)  # 九号字


def setup_page(doc):
    """设置页面布局"""
    section = doc.sections[0]

    # 设置页面大小
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT

    # 设置页边距
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    return section


def add_header(section, header_text):
    """添加页眉"""
    header = section.header
    header.is_linked_to_previous = False

    # 清空默认页眉
    for p in header.paragraphs:
        p.clear()

    # 添加页眉内容
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(header_text)
    run.font.name = HEADER_FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), HEADER_FONT_NAME)
    run.font.size = HEADER_FONT_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 添加底部边框线
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_footer(section):
    """添加页脚（页码）"""
    footer = section.footer
    footer.is_linked_to_previous = False

    # 清空默认页脚
    for p in footer.paragraphs:
        p.clear()

    # 添加页脚内容
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加"第 X 页"文本
    run1 = paragraph.add_run("第 ")
    run1.font.name = HEADER_FONT_NAME
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), HEADER_FONT_NAME)
    run1.font.size = HEADER_FONT_SIZE
    run1.font.color.rgb = RGBColor(0, 0, 0)

    # 插入页码域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run2 = paragraph.add_run()
    run2._element.append(fldChar1)
    run2._element.append(instrText)
    run2._element.append(fldChar2)
    run2.font.name = HEADER_FONT_NAME
    run2.font.size = HEADER_FONT_SIZE
    run2.font.color.rgb = RGBColor(0, 0, 0)

    # 添加"页"文本
    run3 = paragraph.add_run(" 页")
    run3.font.name = HEADER_FONT_NAME
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), HEADER_FONT_NAME)
    run3.font.size = HEADER_FONT_SIZE
    run3.font.color.rgb = RGBColor(0, 0, 0)


def read_source_code():
    """读取源代码文件，提取纯代码部分（无空行）"""
    with open(SOURCE_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 去除文件头信息
    code_lines = []
    in_code = False
    for line in lines:
        # 跳过文件头
        if '软件著作权登记申请' in line or '软件名称' in line or '版本号' in line or '开发完成日期' in line:
            continue
        if '前 30 页' in line or '后 30 页' in line:
            in_code = True
            continue
        if line.startswith('===') or line.startswith('---'):
            continue
        if in_code:
            stripped = line.rstrip('\n')
            # 只添加非空行
            if stripped.strip():
                code_lines.append(stripped)

    return code_lines


def extract_code_sections(code_lines):
    """提取前30页和后30页的代码"""
    total_lines = len(code_lines)

    if total_lines >= TOTAL_LINES:
        # 代码足够，取前1500行和后1500行
        front_lines = code_lines[:LINES_PER_PAGE * PAGES_FRONT]
        back_lines = code_lines[-(LINES_PER_PAGE * PAGES_BACK):]
    else:
        # 代码不足，全部使用
        front_lines = code_lines[:min(LINES_PER_PAGE * PAGES_FRONT, total_lines)]
        back_lines = code_lines[-min(LINES_PER_PAGE * PAGES_BACK, total_lines):] if total_lines > LINES_PER_PAGE * PAGES_FRONT else []

    return front_lines, back_lines


def create_source_code_document():
    """创建源代码文档"""
    print("开始生成源代码文档...")

    # 读取源代码
    code_lines = read_source_code()
    print(f"读取到 {len(code_lines)} 行有效代码（无空行）")

    # 截断过长的代码行（确保每行不超过80个字符）
    print("截断过长的代码行...")
    truncated_lines = []
    for line in code_lines:
        if len(line) > 80:
            truncated_lines.append(line[:80])
        else:
            truncated_lines.append(line)

    # 提取前30页和后30页
    front_lines, back_lines = extract_code_sections(truncated_lines)
    print(f"前30页：{len(front_lines)} 行")
    print(f"后30页：{len(back_lines)} 行")

    # 验证行数
    if len(front_lines) != LINES_PER_PAGE * PAGES_FRONT:
        print(f"警告：前30页应有 {LINES_PER_PAGE * PAGES_FRONT} 行，实际 {len(front_lines)} 行")
    if len(back_lines) != LINES_PER_PAGE * PAGES_BACK:
        print(f"警告：后30页应有 {LINES_PER_PAGE * PAGES_BACK} 行，实际 {len(back_lines)} 行")

    # 创建文档
    doc = Document()

    # 设置页面
    section = setup_page(doc)

    # 添加页眉和页脚
    header_text = f"{SOFTWARE_NAME} {VERSION} 源代码文档"
    add_header(section, header_text)
    add_footer(section)

    # 添加文档标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run(f"{SOFTWARE_NAME} {VERSION}")
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle.add_run("软件著作权登记申请 - 源代码文档")
    subtitle_run.font.name = "黑体"
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0, 0, 0)

    # 添加分页
    last_para = doc.paragraphs[-1]
    run = last_para.add_run()
    run.add_break(WD_BREAK.PAGE)

    # 添加前30页代码
    print("添加前30页代码...")
    current_line_num = 1
    line_count = 0

    for i, line in enumerate(front_lines):
        # 确保没有空行
        if not line.strip():
            continue

        # 添加代码行
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(14)  # 固定行距14磅

        # 添加行号和代码（行号占4位，右对齐）
        line_text = f"{current_line_num:4d}  {line}"
        run = paragraph.add_run(line_text)
        run.font.name = CODE_FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = CODE_FONT_SIZE
        run.font.color.rgb = RGBColor(0, 0, 0)

        line_count += 1
        current_line_num += 1

        # 每50行添加分页符（除了最后一页）
        if line_count % LINES_PER_PAGE == 0 and line_count < len(front_lines):
            last_para = doc.paragraphs[-1]
            run = last_para.add_run()
            run.add_break(WD_BREAK.PAGE)

    # 如果有后30页代码，添加分隔页和后30页
    if back_lines:
        # 添加分页
        last_para = doc.paragraphs[-1]
        run = last_para.add_run()
        run.add_break(WD_BREAK.PAGE)

        # 添加分隔标题
        separator_title = doc.add_paragraph()
        separator_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_title.paragraph_format.space_before = Pt(72)  # 空12行
        run = separator_title.add_run("（后 30 页）")
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 添加分页
        last_para = doc.paragraphs[-1]
        run = last_para.add_run()
        run.add_break(WD_BREAK.PAGE)

        # 添加后30页代码
        print("添加后30页代码...")
        line_count = 0

        for i, line in enumerate(back_lines):
            # 确保没有空行
            if not line.strip():
                continue

            # 添加代码行
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(14)  # 固定行距14磅

            # 添加行号和代码
            line_text = f"{current_line_num:4d}  {line}"
            run = paragraph.add_run(line_text)
            run.font.name = CODE_FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = CODE_FONT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)

            line_count += 1
            current_line_num += 1

            # 每50行添加分页符（除了最后一页）
            if line_count % LINES_PER_PAGE == 0 and line_count < len(back_lines):
                last_para = doc.paragraphs[-1]
                run = last_para.add_run()
                run.add_break(WD_BREAK.PAGE)

    # 保存文档
    doc.save(OUTPUT_FILE)
    print(f"\n文档已保存到: {OUTPUT_FILE}")
    print(f"文件大小: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")

    return True


def main():
    """主函数"""
    try:
        success = create_source_code_document()
        if success:
            print("\n✅ 源代码文档生成完成！")
        else:
            print("\n❌ 文档生成失败")
            return 1
    except Exception as e:
        print(f"\n❌ 生成失败: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == "__main__":
    exit(main())
