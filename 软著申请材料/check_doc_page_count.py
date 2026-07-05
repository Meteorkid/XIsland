#!/usr/bin/env python3
"""
检查文档实际页数
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def check_doc_page_count():
    """检查文档实际页数"""
    print("检查文档实际页数...")

    doc = Document(OUTPUT_FILE)

    # 检查页面设置
    section = doc.sections[0]

    # 检查段落格式
    print(f"\n=== 段落格式检查 ===")
    para_count = 0
    for i, para in enumerate(doc.paragraphs):
        # 检查行距
        if para.paragraph_format.line_spacing:
            line_spacing = para.paragraph_format.line_spacing
            if hasattr(line_spacing, 'pt'):
                if para_count < 5:  # 只显示前5个
                    print(f"段落 {i}: 行距 {line_spacing.pt} pt")
            para_count += 1

    # 检查字体大小
    print(f"\n=== 字体大小检查 ===")
    font_count = 0
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run.font.size:
                if font_count < 5:  # 只显示前5个
                    print(f"段落 {i}: 字体大小 {run.font.size.pt} pt")
                font_count += 1
                break

    # 检查分页符位置
    print(f"\n=== 分页符位置 ===")
    page_break_positions = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run._element.xml.find('w:br') != -1 and 'type="page"' in run._element.xml:
                page_break_positions.append(i)
                break

    print(f"分页符数量: {len(page_break_positions)}")
    print(f"前5个分页符位置: {page_break_positions[:5]}")
    print(f"后5个分页符位置: {page_break_positions[-5:]}")

    # 检查每个分页符前的代码行数
    print(f"\n=== 每个分页符前的代码行数 ===")
    code_count = 0
    for i, pos in enumerate(page_break_positions[:10]):  # 只检查前10个
        # 统计前一个分页符到当前分页符之间的代码行数
        if i == 0:
            start = 0
        else:
            start = page_break_positions[i-1] + 1

        count = 0
        for j in range(start, pos):
            text = doc.paragraphs[j].text.strip()
            if text and text[0].isdigit() and '  ' in text:
                count += 1

        print(f"分页符 {i+1} (位置 {pos}): 前 {count} 行代码")


if __name__ == "__main__":
    check_doc_page_count()
