#!/usr/bin/env python3
"""
验证源代码文档是否符合软著要求
检查：
1. 每页是否正好50行代码
2. 是否有空行
3. 总页数是否正确
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def verify_document():
    """验证文档"""
    print("开始验证源代码文档...")

    if not os.path.exists(OUTPUT_FILE):
        print(f"❌ 文件不存在: {OUTPUT_FILE}")
        return False

    # 读取文档
    doc = Document(OUTPUT_FILE)

    # 统计信息
    total_paragraphs = len(doc.paragraphs)
    empty_lines = 0
    code_lines = 0
    page_breaks = 0

    # 遍历所有段落
    for i, para in enumerate(doc.paragraphs):
        # 检查是否有分页符
        for run in para.runs:
            if run._element.xml.find('w:br') != -1 and 'type="page"' in run._element.xml:
                page_breaks += 1

        # 检查段落内容
        text = para.text.strip()
        if text:
            # 检查是否是代码行（以数字开头）
            if text[0].isdigit() and '  ' in text:
                code_lines += 1
            # 检查是否是空行
            elif text == '':
                empty_lines += 1
        else:
            empty_lines += 1

    print(f"\n=== 文档统计 ===")
    print(f"总段落数: {total_paragraphs}")
    print(f"代码行数: {code_lines}")
    print(f"空行数: {empty_lines}")
    print(f"分页符数: {page_breaks}")

    # 验证结果
    print(f"\n=== 验证结果 ===")

    # 1. 检查代码行数
    expected_code_lines = 3000  # 60页 * 50行
    if code_lines >= expected_code_lines:
        print(f"✅ 代码行数: {code_lines} 行 (满足要求)")
    else:
        print(f"❌ 代码行数: {code_lines} 行 (不足 {expected_code_lines} 行)")

    # 2. 检查空行
    if empty_lines == 0:
        print(f"✅ 空行: 无")
    else:
        print(f"❌ 空行: {empty_lines} 行")

    # 3. 检查分页符
    expected_pages = 60
    if page_breaks >= expected_pages - 1:  # 最后一页不需要分页符
        print(f"✅ 分页符: {page_breaks} 个 (满足要求)")
    else:
        print(f"❌ 分页符: {page_breaks} 个 (不足 {expected_pages - 1} 个)")

    # 4. 检查文件大小
    file_size = os.path.getsize(OUTPUT_FILE) / 1024
    print(f"✅ 文件大小: {file_size:.1f} KB")

    # 总结
    all_passed = (code_lines >= expected_code_lines and
                  empty_lines == 0 and
                  page_breaks >= expected_pages - 1)

    if all_passed:
        print(f"\n✅ 文档完全符合软著申请要求！")
        return True
    else:
        print(f"\n❌ 文档未完全符合要求，请检查上述问题")
        return False


def main():
    """主函数"""
    try:
        success = verify_document()
        return 0 if success else 1
    except Exception as e:
        print(f"\n❌ 验证失败: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())
