#!/usr/bin/env python3
"""
检查文档实际页数和内容
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def check_actual_pages():
    """检查实际页数"""
    print("检查文档实际页数...")

    doc = Document(OUTPUT_FILE)

    # 检查页面设置
    section = doc.sections[0]
    print(f"\n=== 页面设置 ===")
    print(f"页面宽度: {section.page_width.cm:.1f} cm")
    print(f"页面高度: {section.page_height.cm:.1f} cm")
    print(f"左边距: {section.left_margin.cm:.1f} cm")
    print(f"右边距: {section.right_margin.cm:.1f} cm")
    print(f"上边距: {section.top_margin.cm:.1f} cm")
    print(f"下边距: {section.bottom_margin.cm:.1f} cm")

    # 计算可用区域
    usable_width = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    usable_height = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm
    print(f"可用宽度: {usable_width:.1f} cm")
    print(f"可用高度: {usable_height:.1f} cm")

    # 检查段落格式
    print(f"\n=== 段落格式 ===")
    para_count = 0
    for i, para in enumerate(doc.paragraphs[:10]):  # 检查前10个段落
        if para.paragraph_format.line_spacing:
            line_spacing = para.paragraph_format.line_spacing
            if hasattr(line_spacing, 'pt'):
                print(f"段落 {i}: 行距 {line_spacing.pt} pt")
            else:
                print(f"段落 {i}: 行距 {line_spacing}")
        para_count += 1

    # 检查字体大小
    print(f"\n=== 字体大小 ===")
    for i, para in enumerate(doc.paragraphs[:10]):  # 检查前10个段落
        for run in para.runs:
            if run.font.size:
                print(f"段落 {i}: 字体大小 {run.font.size.pt} pt")
                break

    # 估算每页行数
    print(f"\n=== 估算每页行数 ===")
    # 假设字体大小10.5pt，行距14pt
    font_size_pt = 10.5
    line_spacing_pt = 14

    # 计算每页可容纳的行数
    # 1pt = 0.0353 cm
    line_height_cm = line_spacing_pt * 0.0353
    lines_per_page = usable_height / line_height_cm

    print(f"字体大小: {font_size_pt} pt")
    print(f"行距: {line_spacing_pt} pt")
    print(f"行高: {line_height_cm:.2f} cm")
    print(f"每页可容纳行数: {lines_per_page:.1f} 行")

    # 检查分页符数量
    print(f"\n=== 分页符统计 ===")
    page_break_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.xml.find('w:br') != -1 and 'type="page"' in run._element.xml:
                page_break_count += 1
                break

    print(f"分页符数量: {page_break_count}")

    # 检查总段落数
    total_paragraphs = len(doc.paragraphs)
    print(f"总段落数: {total_paragraphs}")


if __name__ == "__main__":
    check_actual_pages()
