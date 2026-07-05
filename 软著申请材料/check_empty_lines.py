#!/usr/bin/env python3
"""
检查文档中的空行
"""

from docx import Document
import os

OUTPUT_FILE = "/Users/meteor/github/xisland/软著申请材料/XIsland_源代码文档.docx"


def check_empty_lines():
    """检查空行"""
    print("检查文档中的空行...")

    doc = Document(OUTPUT_FILE)

    # 统计信息
    total_paragraphs = len(doc.paragraphs)
    empty_lines = []
    non_empty_lines = []

    # 遍历所有段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            empty_lines.append(i)
        else:
            non_empty_lines.append(i)

    print(f"\n=== 文档统计 ===")
    print(f"总段落数: {total_paragraphs}")
    print(f"非空段落数: {len(non_empty_lines)}")
    print(f"空段落数: {len(empty_lines)}")

    # 显示空行位置
    if empty_lines:
        print(f"\n=== 空行位置 ===")
        for pos in empty_lines[:20]:  # 只显示前20个
            print(f"段落 {pos}")
        if len(empty_lines) > 20:
            print(f"... 还有 {len(empty_lines) - 20} 个空行")

    # 检查空行是否在代码区域
    print(f"\n=== 空行分析 ===")
    code_start = None
    code_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and text[0].isdigit() and '  ' in text:
            if code_start is None:
                code_start = i
            code_end = i

    if code_start and code_end:
        print(f"代码区域: 段落 {code_start} 到 {code_end}")

        # 检查代码区域内的空行
        code_empty_lines = [pos for pos in empty_lines if code_start <= pos <= code_end]
        if code_empty_lines:
            print(f"代码区域内空行: {len(code_empty_lines)} 个")
            for pos in code_empty_lines[:10]:
                print(f"  - 段落 {pos}")
        else:
            print(f"代码区域内无空行 ✓")

        # 检查代码区域外的空行
        non_code_empty_lines = [pos for pos in empty_lines if pos < code_start or pos > code_end]
        if non_code_empty_lines:
            print(f"代码区域外空行: {len(non_code_empty_lines)} 个")
            for pos in non_code_empty_lines:
                print(f"  - 段落 {pos} (标题/分隔区域)")
    else:
        print(f"未找到代码区域")

    return len(empty_lines) == 0 or (code_start and code_end and len([pos for pos in empty_lines if code_start <= pos <= code_end]) == 0)


def main():
    """主函数"""
    try:
        success = check_empty_lines()
        if success:
            print(f"\n✅ 文档代码区域无空行！")
        else:
            print(f"\n❌ 文档代码区域有空行，请检查")
        return 0
    except Exception as e:
        print(f"\n❌ 检查失败: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())
